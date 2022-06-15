from pprint import pprint
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


res = requests.get(
    "https://www.professormesser.com/free-a-plus-training/220-1001/220-1000-training-course/")
contents = res.text
soup = BeautifulSoup(contents, "html.parser")
document = Document()

all = soup.select(
    "h1, .entry-content h1, h2, .entry-content h2, h3, .entry-content h3, h4, .entry-content h4, h5, .entry-content h5, h6, .entry-content h6, .site-title, .site-title a",
    limit=220)

heading = document.add_heading(all[0].getText(), 0)
heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
for a in all:
    if "Section" in a.getText():
        if "CompTIA" in a.getText():
            text = a.getText().split("Section")[0]
            print(text)
            subheading1 = document.add_heading(text, 1)
        else:
            text = a.getText()
            print(text)
            subheading2 = document.add_heading(text, 2)
            subheading2.paragraph_format.left_indent = Inches(0.3)
    else:
        if a.name == "h3":
            print(f"Subsection: {a.getText()}")
            subsection = document.add_heading(a.getText(), 3)
            subsection.paragraph_format.left_indent = Inches(0.5)
        elif a.name == "h4":
            print(f"Video Title: {a.getText()}")
            video_title = document.add_heading(a.getText(), 4)
            video_title.paragraph_format.left_indent = Inches(0.9)


        
        
    
document.save('free-a-plus-training.docx')
